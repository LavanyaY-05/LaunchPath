# import os
# import io
# from typing import Optional, List
# from fastapi import FastAPI, File, UploadFile, HTTPException, Query
# from fastapi.middleware.cors import CORSMiddleware
# from pydantic import BaseModel
# from dotenv import load_dotenv

# from graph import app_graph
# from retrieval import search

# load_dotenv()

# app = FastAPI(
#     title="LaunchPath Backend API",
#     description="AI Advisor for Early-Stage Entrepreneurship (FastAPI + LangGraph + Supabase)",
#     version="1.0.0"
# )

# # Enable CORS for Next.js frontend on localhost:3000
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:3000", "http://127.0.0.1:3000", "*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )


# class ChatRequest(BaseModel):
#     message: str
#     domain: Optional[str] = None
#     uploaded_text: Optional[str] = None


# class ChatResponse(BaseModel):
#     answer: str
#     sources: List[str]
#     intent: str
#     follow_ups: List[str]


# class ExploreResponse(BaseModel):
#     answer: str
#     sources: List[str]
#     follow_ups: List[str]


# @app.get("/")
# def read_root():
#     return {"status": "online", "app": "LaunchPath Backend API", "version": "1.0.0"}


# @app.post("/chat", response_model=ChatResponse)
# def chat_endpoint(payload: ChatRequest):
#     if not payload.message or not payload.message.strip():
#         raise HTTPException(status_code=400, detail="Message cannot be empty.")

#     initial_state = {
#         "message": payload.message.strip(),
#         "uploaded_text": payload.uploaded_text.strip() if payload.uploaded_text else None,
#         "domain": payload.domain.strip() if payload.domain else None,
#         "intent": "general_qa",
#         "retrieved_chunks": [],
#         "relevance_ok": False,
#         "answer": "",
#         "sources": [],
#         "follow_ups": []
#     }

#     try:
#         final_state = app_graph.invoke(initial_state)
#         return ChatResponse(
#             answer=final_state.get("answer", "I don't have relevant information on that right now."),
#             sources=final_state.get("sources", []),
#             intent=final_state.get("intent", "general_qa"),
#             follow_ups=final_state.get("follow_ups", [])
#         )
#     except Exception as e:
#         print(f"[Error in /chat]: {e}")
#         raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


# @app.get("/explore/{domain}", response_model=ExploreResponse)
# def explore_domain_endpoint(domain: str):
#     clean_domain = domain.strip().lower()
    
#     domain_queries = {
#         "freelancing": "freelancing client approach pricing GenAI developer UI UX designer web developer content writer",
#         "schemes": "startup india seed fund samridh meity dpiit recognition udyam msme government schemes",
#         "investors": "outreach norms private funding alternatives investor expectations pitch deck essentials",
#         "local_business": "bakery marketing customer retention local delivery logistics small business",
#         "failures": "laundry startup postmortem smart waste tech failure patterns starsky robotics postmortem",
#         "roadmap": "freelance designer roadmap freelance developer roadmap GenAI developer roadmap skills",
#         "pitch_deck": "what a strong early-stage pitch deck should contain portfolio guidance essentials"
#     }

#     query = domain_queries.get(clean_domain, f"{clean_domain} guidance roadmap schemes investors")

#     initial_state = {
#         "message": f"Provide comprehensive structured exploration guide for {clean_domain}",
#         "uploaded_text": None,
#         "domain": clean_domain if clean_domain in domain_queries else None,
#         "intent": "general_qa",
#         "retrieved_chunks": [],
#         "relevance_ok": False,
#         "answer": "",
#         "sources": [],
#         "follow_ups": []
#     }

#     try:
#         final_state = app_graph.invoke(initial_state)
#         return ExploreResponse(
#             answer=final_state.get("answer", f"No structured guide found for {domain}."),
#             sources=final_state.get("sources", []),
#             follow_ups=final_state.get("follow_ups", [])
#         )
#     except Exception as e:
#         print(f"[Error in /explore]: {e}")
#         raise HTTPException(status_code=500, detail=f"Failed to generate explore content: {str(e)}")


# @app.post("/upload")
# async def upload_file_endpoint(file: UploadFile = File(...)):
#     # Check max file size (3MB)
#     contents = await file.read()
#     if len(contents) > 3 * 1024 * 1024:
#         raise HTTPException(status_code=400, detail="File size exceeds maximum limit of 3MB.")

#     filename = file.filename or "uploaded_document"
#     ext = os.path.splitext(filename)[1].lower()

#     extracted_text = ""

#     try:
#         if ext == ".txt":
#             extracted_text = contents.decode("utf-8", errors="ignore").strip()
#         elif ext == ".pdf":
#             try:
#                 import pypdf
#                 reader = pypdf.PdfReader(io.BytesIO(contents))
#                 pages_text = [p.extract_text() for p in reader.pages if p.extract_text()]
#                 extracted_text = "\n".join(pages_text).strip()
#             except Exception as pe:
#                 raise HTTPException(status_code=400, detail=f"Could not parse PDF file: {pe}")
#         elif ext in [".docx", ".doc"]:
#             try:
#                 import docx
#                 doc = docx.Document(io.BytesIO(contents))
#                 paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
#                 extracted_text = "\n".join(paragraphs).strip()
#             except Exception as de:
#                 raise HTTPException(status_code=400, detail=f"Could not parse DOCX file: {de}")
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file format. Please upload .pdf, .docx, or .txt")

#         if not extracted_text:
#             raise HTTPException(status_code=400, detail="Extracted text is empty or could not be read.")

#         return {
#             "extracted_text": extracted_text,
#             "filename": filename
#         }
#     except HTTPException:
#         raise
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Failed to process uploaded file: {str(e)}")


import os
import io
from typing import Optional, List
from fastapi import FastAPI, File, UploadFile, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv

from graph import app_graph
from retrieval import search

load_dotenv()

app = FastAPI(
    title="LaunchPath Backend API",
    description="AI Advisor for Early-Stage Entrepreneurship (FastAPI + LangGraph + Supabase)",
    version="1.0.0"
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "https://launch-path-beryl.vercel.app"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)



class ChatRequest(BaseModel):
    message: str
    domain: Optional[str] = None
    uploaded_files: Optional[List[dict]] = None
    previous_domain: Optional[str] = None
    previous_role: Optional[str] = None


class ChatResponse(BaseModel):
    answer: str
    sources: List[str]
    intent: str
    follow_ups: List[str]
    domain: Optional[str] = None
    role: Optional[str] = None

class ExploreResponse(BaseModel):
    answer: str
    sources: List[str]
    follow_ups: List[str]


@app.get("/")
def read_root():
    return {"status": "online", "app": "LaunchPath Backend API", "version": "1.0.0"}


@app.post("/chat", response_model=ChatResponse)
def chat_endpoint(payload: ChatRequest):
    if not payload.message or not payload.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty.")

    uploaded_text = None
    if payload.uploaded_files and isinstance(payload.uploaded_files, list):
        file_texts = [f.get("extracted_text", "") for f in payload.uploaded_files if isinstance(f, dict) and f.get("extracted_text")]
        uploaded_text = "\n\n".join(file_texts).strip() if file_texts else None

    initial_state = {
        "message": payload.message.strip(),
        "uploaded_text": uploaded_text,
        "uploaded_files": payload.uploaded_files if isinstance(payload.uploaded_files, list) else [],
        "domain": payload.domain.strip() if payload.domain else None,
        "previous_domain": payload.previous_domain.strip() if payload.previous_domain else None,
        "intent": "general_qa",
        "retrieved_chunks": [],
        "relevance_ok": False,
        "answer": "",
        "sources": [],
        "follow_ups": []
    }

    try:
        final_state = app_graph.invoke(initial_state)
        return ChatResponse(
            answer=final_state.get("answer", "I don't have relevant information on that right now."),
            sources=final_state.get("sources", []),
            intent=final_state.get("intent", "general_qa"),
            follow_ups=final_state.get("follow_ups", []),
            domain=final_state.get("domain"),
            role=final_state.get("role")
        )
    except Exception as e:
        print(f"[Error in /chat]: {e}")
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


@app.get("/explore/{domain}", response_model=ExploreResponse)
def explore_domain_endpoint(domain: str):
    clean_domain = domain.strip().lower()

    domain_queries = {
        "freelancing": "freelancing client approach pricing GenAI developer web developer",
        "schemes": "startup india seed fund samridh meity dpiit recognition government schemes",
        "investors": "outreach norms private funding alternatives investor expectations pitch deck essentials",
        "startups": "genai customer support local delivery logistics startup ideas howto",
        "roadmap": "freelance developer roadmap GenAI developer roadmap skills",
        "pitch_deck": "what a strong early-stage pitch deck should contain portfolio guidance essentials"
    }

    query = domain_queries.get(clean_domain, f"{clean_domain} guidance roadmap schemes investors")

    initial_state = {
        "message": f"Provide comprehensive structured exploration guide for {clean_domain}",
        "uploaded_text": None,
        "domain": clean_domain if clean_domain in domain_queries else None,
        "previous_domain": None,   # NEW: explicit, explore never carries chat context
        "intent": "general_qa",
        "retrieved_chunks": [],
        "relevance_ok": False,
        "answer": "",
        "sources": [],
        "follow_ups": []
    }

    try:
        final_state = app_graph.invoke(initial_state)
        return ExploreResponse(
            answer=final_state.get("answer", f"No structured guide found for {domain}."),
            sources=final_state.get("sources", []),
            follow_ups=final_state.get("follow_ups", [])
        )
    except Exception as e:
        print(f"[Error in /explore]: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate explore content: {str(e)}")


@app.post("/upload")
async def upload_file_endpoint(file: UploadFile = File(...)):
    contents = await file.read()
    if len(contents) > 3 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size exceeds maximum limit of 3MB.")

    filename = file.filename or "uploaded_document"
    ext = os.path.splitext(filename)[1].lower()

    extracted_text = ""

    try:
        if ext == ".txt":
            extracted_text = contents.decode("utf-8", errors="ignore").strip()
        elif ext == ".pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(contents))
                pages_text = [p.extract_text() for p in reader.pages if p.extract_text()]
                extracted_text = "\n".join(pages_text).strip()
            except Exception as pe:
                raise HTTPException(status_code=400, detail=f"Could not parse PDF file: {pe}")
        elif ext in [".docx", ".doc"]:
            try:
                import docx
                doc = docx.Document(io.BytesIO(contents))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                extracted_text = "\n".join(paragraphs).strip()
            except Exception as de:
                raise HTTPException(status_code=400, detail=f"Could not parse DOCX file: {de}")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload .pdf, .docx, or .txt")

        if not extracted_text:
            raise HTTPException(status_code=400, detail="Extracted text is empty or could not be read.")

        return {
            "extracted_text": extracted_text,
            "filename": filename
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process uploaded file: {str(e)}")
